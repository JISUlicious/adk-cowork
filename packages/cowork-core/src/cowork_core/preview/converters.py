"""Convert workspace files into previewable payloads.

Each converter takes a ``Path`` and returns a ``PreviewResult`` with the
converted body (bytes) and its content type.  The main entry point
``preview_file`` dispatches on file extension.

Caching by content hash is handled in ``cache.py``.
"""

from __future__ import annotations

import csv
import hashlib
import html
import io
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any


@dataclass(frozen=True)
class PreviewResult:
    body: bytes
    content_type: str
    content_hash: str


def content_hash(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()[:16]


_IMAGE_TYPES: dict[str, str] = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".svg": "image/svg+xml",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
    ".ico": "image/x-icon",
}

# Extensions that get the plain-text preview fallback (rendered as a
# scrollable ``<pre>``). Covers the files users typically drop into a
# workspace: logs, configs, scripts, and most source languages. If the
# extension is missing here but the content looks textual (no NUL in
# the first 4 KB), we still fall through to the text preview.
_TEXT_EXTENSIONS = {
    # Plain text / prose
    ".txt", ".log", ".rst", ".adoc", ".asciidoc", ".org", ".tex",
    # Configs
    ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf",
    ".env", ".properties", ".editorconfig",
    ".gitignore", ".gitattributes", ".dockerignore",
    # Shell
    ".sh", ".bash", ".zsh", ".fish", ".ps1",
    # Source
    ".py", ".pyi", ".pyx", ".pyw",
    ".js", ".jsx", ".ts", ".tsx", ".mjs", ".cjs",
    ".html", ".htm", ".css", ".scss", ".sass", ".less",
    ".xml", ".xhtml",
    ".c", ".h", ".cpp", ".hpp", ".cc", ".hh", ".cxx", ".ipp",
    ".rs", ".go", ".java", ".kt", ".kts", ".scala",
    ".swift", ".m", ".mm", ".rb", ".php", ".pl", ".lua", ".r", ".jl",
    ".sql", ".graphql", ".gql", ".proto",
    ".tf", ".hcl", ".nix",
    ".vim", ".el", ".lisp", ".clj", ".cljs", ".ex", ".exs",
    ".patch", ".diff",
    ".makefile", ".mk",
}

# Exact filenames (no extension, or extension that would trip a
# different dispatch rule) that should still render as text.
_TEXT_FILENAMES = {
    "Dockerfile",
    "Containerfile",
    "Makefile",
    "makefile",
    "GNUmakefile",
    "CMakeLists.txt",
    "Rakefile",
    "Gemfile",
    "Gemfile.lock",
    "Pipfile",
    "Pipfile.lock",
    "requirements.txt",
    "Procfile",
    "LICENSE",
    "NOTICE",
    "README",
    "CHANGELOG",
    "AUTHORS",
    "CONTRIBUTORS",
    ".env",
    ".gitignore",
    ".dockerignore",
}

_SUPPORTED_EXTENSIONS = (
    {".md", ".docx", ".pdf", ".xlsx", ".csv"}
    | set(_IMAGE_TYPES)
    | _TEXT_EXTENSIONS
)


def _looks_like_text(path: Path, sniff_bytes: int = 4096) -> bool:
    """Best-effort heuristic: no NUL byte in the first few KB → text.

    Misses some edge cases (UTF-16 files with NUL-padded ASCII will be
    treated as binary and rejected; rare). Fine for v1 — the whitelist
    above already covers everything common, and the NUL sniff just
    catches the long tail without producing garbage for real binaries.
    """
    try:
        with path.open("rb") as f:
            chunk = f.read(sniff_bytes)
    except OSError:
        return False
    return b"\x00" not in chunk


def preview_file(path: Path) -> PreviewResult:
    if not path.exists():
        raise FileNotFoundError(f"file not found: {path}")

    ext = path.suffix.lower()
    chash = content_hash(path)

    if ext in _IMAGE_TYPES:
        return PreviewResult(
            body=path.read_bytes(),
            content_type=_IMAGE_TYPES[ext],
            content_hash=chash,
        )

    converter = _CONVERTERS.get(ext)
    if converter is not None:
        return converter(path, chash)

    # Text fallback: known extensions, known filenames, or a NUL-sniff
    # for unclassified files. Catches .py / .txt / .log / .json / etc.
    if (
        ext in _TEXT_EXTENSIONS
        or path.name in _TEXT_FILENAMES
        or _looks_like_text(path)
    ):
        return _preview_text(path, chash)

    raise ValueError(
        f"unsupported preview format: {ext or path.name!r} "
        f"(supported: {', '.join(sorted(_SUPPORTED_EXTENSIONS))})"
    )


_TEXT_PREVIEW_MAX_BYTES = 500_000


def _preview_text(path: Path, chash: str) -> PreviewResult:
    """Render an arbitrary text file inside a themed HTML ``<pre>``.

    Truncates at 500 KB so a 50 MB log doesn't blow the iframe; the
    agent's own fs tools have their own (larger) caps for actual reads.
    Decodes with ``errors='replace'`` so the preview always renders
    something, even for files that sniff as text but aren't valid UTF-8.
    """
    data = path.read_bytes()
    truncated = len(data) > _TEXT_PREVIEW_MAX_BYTES
    if truncated:
        data = data[:_TEXT_PREVIEW_MAX_BYTES]
    text = data.decode("utf-8", errors="replace")
    if truncated:
        text += (
            f"\n\n… [truncated at {_TEXT_PREVIEW_MAX_BYTES:,} bytes]"
        )
    body = _wrap_text(text, path.name)
    return PreviewResult(
        body=body.encode("utf-8"),
        content_type="text/html",
        content_hash=chash,
    )


def _wrap_text(text: str, title: str) -> str:
    safe_title = html.escape(title)
    safe_text = html.escape(text)
    style = (
        ":root{color-scheme:light dark}"
        "body{font-family:ui-monospace,'SF Mono',Menlo,Consolas,monospace;"
        "font-size:12px;line-height:1.5;margin:0;padding:1em;"
        "background:#fff;color:#111}"
        "pre{white-space:pre-wrap;word-break:break-word;margin:0}"
        "@media (prefers-color-scheme: dark){"
        "body{background:#0b0f17;color:#e5e7eb}"
        "}"
    )
    return (
        f"<!DOCTYPE html><html><head><meta charset='utf-8'>"
        f"<title>{safe_title}</title>"
        f"<meta name='color-scheme' content='light dark'>"
        f"<style>{style}</style>"
        f"</head><body><pre>{safe_text}</pre></body></html>"
    )


def _preview_md(path: Path, chash: str) -> PreviewResult:
    from markdown_it import MarkdownIt

    md = MarkdownIt("commonmark", {"html": False})
    source = path.read_text(encoding="utf-8")
    rendered = md.render(source)
    body = _wrap_html(rendered, path.name)
    return PreviewResult(body=body.encode("utf-8"), content_type="text/html", content_hash=chash)


def _wrap_html(inner: str, title: str) -> str:
    safe_title = html.escape(title)
    style = (
        ":root{color-scheme:light dark}"
        "body{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;"
        "max-width:48em;margin:2em auto;padding:0 1em;"
        "background:#fff;color:#111}"
        "pre{background:#f4f4f4;padding:1em;overflow-x:auto}"
        "code{background:#f4f4f4;padding:0.2em 0.4em}"
        "a{color:#2563eb}"
        "@media (prefers-color-scheme: dark){"
        "body{background:#0b0f17;color:#e5e7eb}"
        "pre,code{background:#1f2937;color:#e5e7eb}"
        "a{color:#60a5fa}"
        "blockquote{border-left:3px solid #374151;color:#9ca3af}"
        "hr{border-color:#374151}"
        "th,td{border-color:#374151}"
        "}"
    )
    return (
        f"<!DOCTYPE html><html><head><meta charset='utf-8'>"
        f"<title>{safe_title}</title>"
        f"<meta name='color-scheme' content='light dark'>"
        f"<style>{style}</style>"
        f"</head><body>{inner}</body></html>"
    )


def _preview_docx(path: Path, chash: str) -> PreviewResult:
    from docx import Document

    doc = Document(str(path))
    paragraphs: list[dict[str, Any]] = []
    for para in doc.paragraphs:
        paragraphs.append({
            "text": para.text,
            "style": para.style.name if para.style else None,
        })

    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append(rows)

    payload: dict[str, Any] = {"paragraphs": paragraphs}
    if tables:
        payload["tables"] = tables

    body = json.dumps(payload, ensure_ascii=False, indent=2)
    return PreviewResult(body=body.encode("utf-8"), content_type="application/json", content_hash=chash)


def _preview_pdf(path: Path, chash: str) -> PreviewResult:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    meta = reader.metadata
    pages: list[dict[str, Any]] = []
    for i, page in enumerate(reader.pages):
        pages.append({
            "page": i + 1,
            "text": page.extract_text() or "",
        })

    payload: dict[str, Any] = {
        "page_count": len(reader.pages),
        "metadata": {
            "title": meta.title if meta else None,
            "author": meta.author if meta else None,
            "subject": meta.subject if meta else None,
            "creator": meta.creator if meta else None,
        },
        "pages": pages,
    }
    body = json.dumps(payload, ensure_ascii=False, indent=2)
    return PreviewResult(body=body.encode("utf-8"), content_type="application/json", content_hash=chash)


def _preview_xlsx(path: Path, chash: str) -> PreviewResult:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    sheets: list[dict[str, Any]] = []
    for name in wb.sheetnames:
        ws = wb[name]
        rows: list[list[Any]] = []
        for row in ws.iter_rows(values_only=True):
            rows.append([_cell_value(c) for c in row])
        schema: list[str] = []
        if rows:
            schema = [str(c) if c is not None else f"col_{i}" for i, c in enumerate(rows[0])]
        sheets.append({
            "name": name,
            "schema": schema,
            "rows": rows,
        })
    wb.close()

    body = json.dumps(sheets, ensure_ascii=False, indent=2, default=str)
    return PreviewResult(body=body.encode("utf-8"), content_type="application/json", content_hash=chash)


def _preview_csv(path: Path, chash: str) -> PreviewResult:
    text = path.read_text(encoding="utf-8")
    dialect = csv.Sniffer().sniff(text[:4096]) if len(text) > 0 else csv.excel
    reader = csv.reader(io.StringIO(text), dialect)
    all_rows = list(reader)

    schema: list[str] = []
    data_rows: list[list[str]] = []
    if all_rows:
        schema = all_rows[0]
        data_rows = all_rows[1:]

    payload = {"schema": schema, "rows": data_rows}
    body = json.dumps(payload, ensure_ascii=False, indent=2)
    return PreviewResult(body=body.encode("utf-8"), content_type="application/json", content_hash=chash)


def _cell_value(v: Any) -> Any:
    if v is None:
        return None
    if isinstance(v, (int, float, bool)):
        return v
    return str(v)


_CONVERTERS: dict[str, Any] = {
    ".md": _preview_md,
    ".docx": _preview_docx,
    ".pdf": _preview_pdf,
    ".xlsx": _preview_xlsx,
    ".csv": _preview_csv,
}
