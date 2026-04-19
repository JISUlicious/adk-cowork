"""M1.9 tests for preview converters and cache."""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from cowork_core.preview import preview_file
from cowork_core.preview.cache import PreviewCache
from cowork_core.preview.converters import content_hash


@pytest.fixture()
def tmp(tmp_path: Path) -> Path:
    return tmp_path


class TestContentHash:
    def test_deterministic(self, tmp: Path) -> None:
        f = tmp / "a.txt"
        f.write_text("hello")
        assert content_hash(f) == content_hash(f)

    def test_changes_with_content(self, tmp: Path) -> None:
        f = tmp / "a.txt"
        f.write_text("v1")
        h1 = content_hash(f)
        f.write_text("v2")
        h2 = content_hash(f)
        assert h1 != h2


class TestMarkdownPreview:
    def test_converts_to_html(self, tmp: Path) -> None:
        f = tmp / "readme.md"
        f.write_text("# Hello\n\nWorld")
        result = preview_file(f)
        assert result.content_type == "text/html"
        assert b"<h1>Hello</h1>" in result.body
        assert b"World" in result.body

    def test_html_injection_blocked(self, tmp: Path) -> None:
        f = tmp / "bad.md"
        f.write_text("<script>alert('xss')</script>")
        result = preview_file(f)
        assert b"<script>" not in result.body


class TestDocxPreview:
    def test_extracts_paragraphs(self, tmp: Path) -> None:
        from docx import Document

        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Body text here.")
        path = tmp / "test.docx"
        doc.save(str(path))

        result = preview_file(path)
        assert result.content_type == "application/json"
        data = json.loads(result.body)
        assert "paragraphs" in data
        texts = [p["text"] for p in data["paragraphs"]]
        assert "Title" in texts
        assert "Body text here." in texts

    def test_extracts_tables(self, tmp: Path) -> None:
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "2"
        path = tmp / "table.docx"
        doc.save(str(path))

        result = preview_file(path)
        data = json.loads(result.body)
        assert "tables" in data
        assert data["tables"][0][0] == ["A", "B"]


class TestPdfPreview:
    def test_extracts_text(self, tmp: Path) -> None:
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        path = tmp / "test.pdf"
        with path.open("wb") as f:
            writer.write(f)

        result = preview_file(path)
        assert result.content_type == "application/json"
        data = json.loads(result.body)
        assert data["page_count"] == 1
        assert "metadata" in data
        assert "pages" in data


class TestXlsxPreview:
    def test_extracts_rows(self, tmp: Path) -> None:
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        assert ws is not None
        ws.append(["Name", "Value"])
        ws.append(["Alice", 42])
        ws.append(["Bob", 99])
        path = tmp / "test.xlsx"
        wb.save(str(path))

        result = preview_file(path)
        assert result.content_type == "application/json"
        data = json.loads(result.body)
        assert len(data) == 1
        sheet = data[0]
        assert sheet["schema"] == ["Name", "Value"]
        assert len(sheet["rows"]) == 3


class TestCsvPreview:
    def test_extracts_schema_and_rows(self, tmp: Path) -> None:
        f = tmp / "data.csv"
        f.write_text("name,age\nAlice,30\nBob,25\n")

        result = preview_file(f)
        assert result.content_type == "application/json"
        data = json.loads(result.body)
        assert data["schema"] == ["name", "age"]
        assert len(data["rows"]) == 2
        assert data["rows"][0] == ["Alice", "30"]

    def test_empty_csv(self, tmp: Path) -> None:
        f = tmp / "empty.csv"
        f.write_text("")
        result = preview_file(f)
        data = json.loads(result.body)
        assert data["schema"] == []
        assert data["rows"] == []


class TestImagePreview:
    def test_png_passthrough(self, tmp: Path) -> None:
        f = tmp / "pixel.png"
        # Minimal 1x1 PNG
        import struct
        import zlib

        def _minimal_png() -> bytes:
            sig = b"\x89PNG\r\n\x1a\n"
            ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
            ihdr = _chunk(b"IHDR", ihdr_data)
            raw = b"\x00\xff\x00\x00"
            idat = _chunk(b"IDAT", zlib.compress(raw))
            iend = _chunk(b"IEND", b"")
            return sig + ihdr + idat + iend

        def _chunk(ctype: bytes, data: bytes) -> bytes:
            c = ctype + data
            return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

        f.write_bytes(_minimal_png())
        result = preview_file(f)
        assert result.content_type == "image/png"
        assert result.body == f.read_bytes()


class TestUnsupportedFormat:
    def test_binary_unknown_ext_rejected(self, tmp: Path) -> None:
        """Unknown extension + binary content (NUL bytes) → reject."""
        f = tmp / "data.xyz"
        f.write_bytes(b"\x00\x01\x02mystery\x00\x03")
        with pytest.raises(ValueError, match="unsupported preview format"):
            preview_file(f)


class TestTextPreview:
    """Text fallback catches .py, .txt, .log, .json, Dockerfiles, and
    unknown-extension files whose content sniffs as text. This is the
    common-files support users expect — without it previews 422 for
    anything that isn't md/docx/pdf/xlsx/csv/image."""

    def test_py_file_renders_as_text(self, tmp: Path) -> None:
        f = tmp / "hello.py"
        f.write_text("print('hi')\n")
        result = preview_file(f)
        assert result.content_type == "text/html"
        # HTML-escaped inside the <pre>; quotes become &#x27;
        assert b"print(" in result.body
        assert b"hi" in result.body

    def test_txt_file_renders(self, tmp: Path) -> None:
        f = tmp / "notes.txt"
        f.write_text("just some notes\n")
        result = preview_file(f)
        assert result.content_type == "text/html"
        assert b"just some notes" in result.body

    def test_html_escaped_in_text_preview(self, tmp: Path) -> None:
        """Source code with angle brackets must not render as HTML."""
        f = tmp / "index.html"
        f.write_text("<script>alert(1)</script>\n")
        result = preview_file(f)
        body = result.body.decode("utf-8")
        assert "&lt;script&gt;" in body
        assert "<script>alert(1)</script>" not in body

    def test_dockerfile_by_name(self, tmp: Path) -> None:
        f = tmp / "Dockerfile"
        f.write_text("FROM python:3.12\n")
        result = preview_file(f)
        assert b"FROM python:3.12" in result.body

    def test_unknown_ext_but_textual_content(self, tmp: Path) -> None:
        """NUL-byte sniff lets .xyz files through if they look like text."""
        f = tmp / "mystery.xyz"
        f.write_text("plain text content\n")
        result = preview_file(f)
        assert b"plain text content" in result.body

    def test_text_preview_truncates_huge_files(self, tmp: Path) -> None:
        f = tmp / "huge.log"
        # Write > 500 KB of text
        f.write_text("x" * 600_000)
        result = preview_file(f)
        assert b"truncated" in result.body


class TestPreviewCache:
    def test_caches_and_returns(self, tmp: Path) -> None:
        cache = PreviewCache(tmp / "cache")
        f = tmp / "test.md"
        f.write_text("# Hi")

        r1 = cache.get(f)
        r2 = cache.get(f)
        assert r1.body == r2.body
        assert r1.content_hash == r2.content_hash

    def test_invalidates_on_change(self, tmp: Path) -> None:
        cache = PreviewCache(tmp / "cache")
        f = tmp / "test.md"
        f.write_text("# V1")
        r1 = cache.get(f)

        f.write_text("# V2")
        r2 = cache.get(f)
        assert r1.content_hash != r2.content_hash
        assert b"V2" in r2.body
